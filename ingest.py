import os
import sys
from langchain_community.document_loaders import PDFPlumberLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document

# --- Configuration ---
DB_DIR = "./chroma_db"
EMBEDDING_MODEL = "BAAI/bge-m3"
SUPPORTED_EXTENSIONS = (".pdf", ".docx")


def load_document(file_path):
    """Extract text from a PDF or DOCX file."""
    print(f"Loading document: {file_path}...")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Error: Could not find '{file_path}'.")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        loader = PDFPlumberLoader(file_path)
        documents = loader.load()
    elif ext == ".docx":
        documents = _load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: '{ext}'. Supported: {SUPPORTED_EXTENSIONS}")
    
    print(f"Successfully loaded {len(documents)} pages/sections.\n")
    return documents


def _load_docx(file_path):
    """Load a DOCX file into LangChain Document objects."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    documents = []

    # Group paragraphs into page-like sections (~3000 chars each)
    current_text = ""
    page_num = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        current_text += text + "\n"

        # Create a new "page" every ~3000 characters (similar to a PDF page)
        if len(current_text) >= 3000:
            page_num += 1
            documents.append(Document(
                page_content=current_text.strip(),
                metadata={"source": file_path, "page": page_num},
            ))
            current_text = ""

    # Don't lose the last section
    if current_text.strip():
        page_num += 1
        documents.append(Document(
            page_content=current_text.strip(),
            metadata={"source": file_path, "page": page_num},
        ))

    # Also extract text from tables
    for table in doc.tables:
        table_text = ""
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_text += " | ".join(row_data) + "\n"
        if table_text.strip():
            page_num += 1
            documents.append(Document(
                page_content=table_text.strip(),
                metadata={"source": file_path, "page": page_num, "type": "table"},
            ))

    return documents

def chunk_document(documents):
    """Step 2: Break the document into smaller semantic chunks."""
    print("Chunking document into smaller pieces...")
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,       # Max characters per chunk
        chunk_overlap=150,    # Overlap to prevent cutting clauses in half
        separators=["\n\n", "\n", " ", ""] # Splits by paragraphs first
    )
    
    chunks = text_splitter.split_documents(documents)
    print(f"Document split into {len(chunks)} chunks.\n")
    return chunks

def create_vector_db(chunks):
    """Step 3 & 4: Convert chunks to vectors and store/append to ChromaDB."""
    print("Loading BGE-M3 embedding model (this may take a minute on first run)...")
    
    embeddings = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': True}
    )
    
    print("Adding documents to Vector Database...")

    # Connect to existing DB or create a new one
    db = Chroma(
        persist_directory=DB_DIR,
        embedding_function=embeddings
    )

    # Add new chunks (appends, does not overwrite existing data)
    db.add_documents(chunks)
    
    print(f"Database successfully updated in '{DB_DIR}' folder!")
    return db

def list_ingested_documents():
    """List all documents currently stored in the database."""
    embeddings = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': True}
    )
    
    if not os.path.exists(DB_DIR):
        print("No database found. Run ingestion first.")
        return []

    db = Chroma(persist_directory=DB_DIR, embedding_function=embeddings)
    collection = db.get()
    
    # Extract unique source filenames from metadata
    sources = set()
    for meta in collection.get("metadatas", []):
        if meta and "source" in meta:
            sources.add(meta["source"])
    
    return sorted(sources)

if __name__ == "__main__":
    # Accept file path as a command-line argument or prompt
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    else:
        file_path = input("Enter the path to the document (PDF or DOCX): ").strip()

    if not file_path:
        print("Error: No file path provided.")
        sys.exit(1)

    print("=== STARTING INGESTION PROCESS ===")
    
    # 1. Load document (PDF or DOCX)
    docs = load_document(file_path)
    
    # 2. Chunk Text
    chunks = chunk_document(docs)
    
    # 3. Embed & Store (appends to existing DB)
    create_vector_db(chunks)
    
    # 4. Show all ingested documents
    print("\n Documents currently in the database:")
    for doc_name in list_ingested_documents():
        print(f"   • {doc_name}")
    
    print("\n=== INGESTION COMPLETE ===")